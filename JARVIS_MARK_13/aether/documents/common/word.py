"""
documents/common/word.py

Handles Word (.docx) document creation, reading, and editing.
Platform-independent business logic using python-docx.
"""

from pathlib import Path
from docx import Document
from aether.documents.common.exceptions import DocumentCorruptedError

def create_word_doc(file_path: Path, content: str | None = None, overwrite: bool = False) -> str:
    """
    Creates a new Word (.docx) document with optional content in the first paragraph.
    """
    if file_path.exists() and not overwrite:
        raise FileExistsError(f"File already exists: {file_path}")
        
    # Create parent directories if missing
    file_path.parent.mkdir(parents=True, exist_ok=True)
    
    try:
        doc = Document()
        if content:
            doc.add_paragraph(content)
        else:
            doc.add_paragraph("")
        doc.save(str(file_path))
        return str(file_path.resolve())
    except Exception as e:
        raise DocumentCorruptedError(f"Failed to create Word document: {e}")

def read_word_doc(file_path: Path) -> str:
    """
    Reads all paragraphs from a Word (.docx) document and returns plain text.
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    if not file_path.is_file():
        raise IsADirectoryError(f"Path is a directory: {file_path}")
        
    try:
        doc = Document(str(file_path))
        text_runs = []
        for paragraph in doc.paragraphs:
            text_runs.append(paragraph.text)
        return "\n".join(text_runs)
    except Exception as e:
        raise DocumentCorruptedError(f"Failed to read Word document structure: {e}")

def edit_word_doc(
    file_path: Path,
    operation: str,
    text: str | None = None,
    old_text: str | None = None,
    new_text: str | None = None
) -> str:
    """
    Modifies an existing Word (.docx) document.
    Operations:
      - 'append': Appends text to the end of the document.
      - 'replace': Replaces all occurrences of old_text with new_text.
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
        
    try:
        doc = Document(str(file_path))
    except Exception as e:
        raise DocumentCorruptedError(f"Failed to load Word document: {e}")
        
    op = operation.lower().strip()
    if op == "append":
        if text is None:
            raise ValueError("Parameter 'text' is required for append operation.")
        doc.add_paragraph(text)
    elif op == "replace":
        if old_text is None or new_text is None:
            raise ValueError("Parameters 'old_text' and 'new_text' are required for replace operation.")
        for paragraph in doc.paragraphs:
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)
    else:
        raise ValueError(f"Unsupported operation '{operation}'. Use 'append' or 'replace'.")
        
    try:
        doc.save(str(file_path))
        return f"Successfully modified document '{file_path.name}' using operation '{op}'."
    except Exception as e:
        raise DocumentCorruptedError(f"Failed to save modified Word document: {e}")
